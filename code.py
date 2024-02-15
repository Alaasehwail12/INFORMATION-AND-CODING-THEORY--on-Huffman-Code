import heapq
import math
import pandas as pd

from docx import Document

doc_path = 'Story.docx'
def build_huffman_tree(freq):
    heap = [[prob, [char, ""]] for char, prob in freq.items()] #for each char has an empty text for the code according to prob.
    heapq.heapify(heap)  #sort heap descending way
    while len(heap) > 1:
        # line 13 + 14 pop the lowest prob. for char in heap
        lo = heapq.heappop(heap)
        hi = heapq.heappop(heap)
        # assign 0 and 1
        for pair in lo[1:]:
            pair[1] = '0' + pair[1]
        for pair in hi[1:]:
            pair[1] = '1' + pair[1]
        heapq.heappush(heap, [lo[0] + hi[0]] + lo[1:] + hi[1:])  #combine char with its codeword
    return heap[0][1:]

def huffman_codes(tree):
    codes = {}
    for char, code in tree:
        codes[char] = code
    return codes


doc = Document(doc_path)
character_counts = {}    #number of ocureences for each char.
characters = set()       #unique char.
total_characters = 0

for paragraph in doc.paragraphs:
    text = paragraph.text.strip().lower()
    for char in text:
        if char.isalpha() or char in ':,.;?!\'"-— ':  #check char.
            characters.add(char)
            character_counts[char] = character_counts.get(char, 0) + 1
            total_characters += 1

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            text = cell.text.strip().lower()
            for char in text:
                if char.isalpha() or char in ':,.;?!\'"-— ':
                    characters.add(char)
                    character_counts[char] = character_counts.get(char, 0) + 1
                    total_characters +=1

unique_characters = sorted(characters) #sort char.


character_probabilities = {}
for char, count in character_counts.items():
    character_probabilities[char] = count / total_characters #calculate the prob.

entropy = 0.0
for char, probability in character_probabilities.items():
    entropy += probability * math.log2(1 / probability)  #calculate the entropy

tree = build_huffman_tree(character_probabilities)

codes = huffman_codes(tree)

bits_per_ascii_character = 8 #ASCII value for each char.
NASCII = total_characters * bits_per_ascii_character #number of bits needed if using ASCII to encode the whole document

print('\n')
print(f"************************ Total Number of Characters in Document:  {total_characters}")

data = {
    'Symbol': unique_characters,
    'Frequency': [character_counts[char] for char in unique_characters],
    'Probability': [character_probabilities[char] for char in unique_characters],
    'Codeword': [codes[char] for char in unique_characters],
    'Length of Codeword': [len(codes[char]) for char in unique_characters]
}

df = pd.DataFrame(data)
print('\n')
print(df.to_string(index=False))
print('\n')

print('************************ Entropy of Characters in the Document: ************************')
print(f'Entropy: {entropy:.6f} bits')
print('\n')

print(f"************************ Number of Bits for ASCII Encoding (N_ASCII): ************************")
print(f"N_ASCII: {NASCII} bits")
print('\n')
#calculate average bits/char.
average_bits_per_character = sum(len(code) * character_probabilities[char] for char, code in codes.items())


print('************************ Average Bits/Character using Huffman Coding: ************************')
print(f'Average Bits/Character: {average_bits_per_character:.6f} bits')
print('\n')
#number of bits needed if using Huffman code to encode the whole document
Nhuffman = sum(len(code) * character_counts[char] for char, code in codes.items())

print('************************ Total Number of Bits Needed To Encode the Entire Story Using Huffman Code (N_Huffman): ************************')
print(f'Total Number of Bits Using Huffman: {Nhuffman} bits')
print('\n')

percentage_Huffman_ASCII = Nhuffman / NASCII * 100
print('************************ Percentage Huffman-ASCII : ************************')
print(f'Percentage of Compression: {percentage_Huffman_ASCII:.6f} ')
print('\n')


